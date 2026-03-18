import sys
import os
import docx

# To install the required library:
# pip install python-docx

def convert_docx_to_text(docx_path):
    """
    Converts a .docx file to a .txt file, preserving Spanish characters
    with UTF-8 encoding without BOM.
    """
    if not os.path.exists(docx_path):
        print(f"Error: File not found at '{docx_path}'")
        return

    try:
        doc = docx.Document(docx_path)
        if doc.tables:
            raise ValueError("El script no soporta archivos con tablas.")
        text = '\n'.join([para.text for para in doc.paragraphs])

        base_name = os.path.splitext(docx_path)[0]
        txt_path = f"{base_name}.txt"

        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)

        print(f"Successfully converted '{docx_path}' to '{txt_path}'")
        return txt_path

    except FileNotFoundError as e:
        print(f"File not found: {e}")
    except PermissionError as e:
        print(f"Permission error: {e}")
    except docx.opc.exceptions.PackageNotFoundError as e:
        print(f"Invalid docx file: {e}")
    except Exception as e:
        print(f"An unexpected error occurred while processing {docx_path}: {e}")
    return None

def process_directory(directory_path):
    """
    Processes all .docx files in a given directory.
    """
    if not os.path.isdir(directory_path):
        print(f"Error: Directory not found at '{directory_path}'")
        return

    print(f"Processing all .docx files in '{directory_path}'...")
    converted_files = []
    for filename in os.listdir(directory_path):
        if filename.lower().endswith(".docx"):
            file_path = os.path.join(directory_path, filename)
            txt_path = convert_docx_to_text(file_path)
            if txt_path:
                converted_files.append(txt_path)
    print("Processing complete.")
    if converted_files:
        print("\nConverted files:")
        for f in converted_files:
            print(f"'{f}'")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python convert_docx_to_text.py <path_to_directory_or_file>")
    else:
        path_arg = sys.argv[1]
        if os.path.isdir(path_arg):
            process_directory(path_arg)
        elif os.path.isfile(path_arg):
            if path_arg.lower().endswith(".docx"):
                convert_docx_to_text(path_arg)
            else:
                print("Error: The specified file is not a .docx file.")
        else:
            print(f"Error: The path '{path_arg}' is not a valid file or directory.")
